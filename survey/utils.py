def get_questions_from_file(test):
    import string
    from .models import Question, Answer
    from docx import Document

    try:
        document = Document(test.file)
    except:
        # TODO proper error handling
        return
    test.questions.delete()
    table = document.tables[0]

    # we definitely should enumerate everything, so it runs pretty slow
    for row_i, row in enumerate(table.rows):
        if row_i == 0:
            pass
        else:
            answers = {}
            answer_i = 0
            question = None

            for p_i, p in enumerate(row.cells[1].paragraphs):
                if p_i == 0:
                    question = Question(
                        test=test, text=p.text, position=row_i)
                    question.save()

                else:
                    if len(p.text.strip()):
                        answer_i += 1
                        answers[answer_i] = {}
                        answers[answer_i]['text'] = p.text.strip()
                        answers[answer_i]['correct'] = False

            alphabet = list(string.ascii_uppercase)
            for char in row.cells[2].text:
                # it can be ints
                try:
                    index = int(char)
                except:
                # or capital letters
                    try:
                        index = alphabet.index(char) + 1
                    except:
                        pass
                try:
                    answers[index]['correct'] = True
                # but a bunch of shit is most common case, so:
                # TODO: proper error handling
                except:
                    print('warning, error in correct answer in question ', row_i, char)

            for key, value in answers.items():
                answer = Answer(question=question,
                                text=value['text'], is_correct=value['correct'])
                answer.save()
